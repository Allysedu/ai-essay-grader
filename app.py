import streamlit as st
import pandas as pd
import pdfplumber
import google.generativeai as genai
import time
import re
import json
import datetime
import os
import base64
import io
import zipfile
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

# --- 🖥️ 앱 기본 설정 (가장 먼저 실행되어야 합니다) ---
st.set_page_config(page_title="AI 에세이 평가 플랫폼", page_icon="🤖", layout="wide")

# --- 🔐 1. 비밀번호 확인 기능 ---
def check_password():
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False
    if not st.session_state.password_correct:
        st.header("🔒 로그인")
        password = st.text_input("비밀번호를 입력하세요", type="password", key="password_input")
        if st.button("로그인", key="login_button"):
            correct_password = st.secrets.get("APP_PASSWORD", "skwlals25") # 🔑 여기에 기본 비밀번호 설정
            if password == correct_password:
                st.session_state.password_correct = True
                st.rerun()
            else:
                st.error("비밀번호가 올바르지 않습니다.")
        return False
    return True

# --- 🧠 보고서 생성 함수 (마크다운) ---
def generate_report_markdown(result_data):
    """학생 한 명의 평가 결과를 화면 표시용 마크다운 텍스트로 만듭니다."""
    parsed_data = result_data.get('평가결과_분석', {})
    report = [
        f"#### 💬 종합 평가",
        f"{parsed_data.get('종합 평가', '내용 없음')}",
        "---",
        f"#### 💯 항목별 상세 평가"
    ]
    itemized_scores = parsed_data.get('항목별 평가', {})
    if itemized_scores:
        for item_name, details in itemized_scores.items():
            report.append(f"**- {item_name} ({details.get('점수', 'N/A')} / {details.get('배점', 'N/A')})**")
            report.append(f"> {details.get('이유', '내용 없음')}")
    else:
        report.append("상세 평가 내용을 불러오지 못했습니다.")
    return "\n".join(report)

# --- 🧠 보고서 생성 함수 (워드 .docx) ---
def generate_report_docx(result_data, eval_name, eval_date):
    """학생 한 명의 평가 결과를 워드(.docx) 파일로 만듭니다."""
    document = docx.Document()
    # 한글 폰트 설정
    style = document.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    file_name = result_data['파일명']
    parsed_data = result_data.get('평가결과_분석', {})
    
    document.add_heading('AI 에세이 평가 상세 분석 보고서', level=1)
    p = document.add_paragraph(); p.add_run('평가명: ').bold = True; p.add_run(eval_name)
    p = document.add_paragraph(); p.add_run('평가일자: ').bold = True; p.add_run(eval_date.strftime('%Y-%m-%d'))
    p = document.add_paragraph(); p.add_run('파일명 (학생): ').bold = True; p.add_run(file_name)
    p = document.add_paragraph(); p.add_run('총점: ').bold = True; p.add_run(f"{parsed_data.get('총점', 'N/A')} 점")
    document.add_heading('💬 종합 평가', level=2)
    document.add_paragraph(parsed_data.get('종합 평가', '내용 없음'))
    document.add_heading('💯 항목별 상세 평가', level=2)
    itemized_scores = parsed_data.get('항목별 평가', {})
    if itemized_scores:
        for item_name, details in itemized_scores.items():
            document.add_heading(f"{item_name} ({details.get('점수', 'N/A')} / {details.get('배점', 'N/A')})", level=3)
            p = document.add_paragraph(); p.add_run('평가 이유: ').bold = True; p.add_run(details.get('이유', '내용 없음'))
    
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- 🖥️ 2. 메인 앱 실행 ---
if check_password():
    # --- AI 초기화 ---
    try:
        if "GOOGLE_AI_API_KEY" in st.secrets:
            genai.configure(api_key=st.secrets["GOOGLE_AI_API_KEY"])
        else:
            st.warning("Google AI API 키가 설정되지 않았습니다. 관리자에게 문의하세요.")
            st.stop()
    except Exception as e:
        st.error(f"Google AI API 키 설정에 실패했습니다: {e}")
        st.stop()

    # --- ✨ 제목 및 프로필 사진 ---
    st.markdown("""<style>.profile-img img {width: 90px; height: 90px; border-radius: 50%; object-fit: cover; margin-top: 10px; margin-bottom: 10px;}</style>""", unsafe_allow_html=True)
    col1, col2 = st.columns([1, 5])
    with col1:
        image_file = 'my_photo.jpg' # ⚠️ 실제 사진 파일 이름으로 변경
        if os.path.exists(image_file):
            with open(image_file, "rb") as f:
                img_base64 = base64.b64encode(f.read()).decode()
                st.markdown(f'<div class="profile-img"><img src="data:image/jpeg;base64,{img_base64}"></div>', unsafe_allow_html=True)
    with col2:
        st.title("AI 에세이 평가 플랫폼")
        st.caption("Ally 교수의 맞춤형 AI 평가 도우미")

    # --- 📂 데이터베이스(JSON 파일) 관리 함수 ---
    HISTORY_FILE = 'evaluation_history.json'
    def load_history():
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                try: return json.load(f)
                except json.JSONDecodeError: return []
        return []
    def save_history(history_data):
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history_data, f, ensure_ascii=False, indent=4)

    history = load_history()
    with st.sidebar:
        st.header("📚 평가 기록 불러오기")
        if history:
            history_options = {f"{item['평가명']} ({item['평가일자']})": i for i, item in enumerate(history)}
            selected_history = st.selectbox("불러올 평가를 선택하세요.", options=list(history_options.keys()))
            if st.button("선택한 평가 기준 불러오기"):
                selected_index = history_options[selected_history]
                st.session_state.criteria_list = history[selected_index]['평가기준']
                st.success(f"'{selected_history}'의 평가 기준을 불러왔습니다.")
                st.rerun()
        else: st.info("저장된 평가 기록이 없습니다.")

    # --- 🧠 AI 응답 분석 함수 ---
    def parse_ai_response(response_text, criteria_list):
        parsed_data = {}
        try:
            summary_match = re.search(r"\[종합 평가\]\s*([\s\S]*?)\s*\[항목별 평가\]", response_text)
            parsed_data['종합 평가'] = summary_match.group(1).strip() if summary_match else "종합 평가 추출 실패"
            scores = {}
            total_score = 0
            for criterion in criteria_list:
                item_name = criterion['항목']
                max_score = criterion['배점']
                pattern = re.compile(rf"- {re.escape(item_name)}:\s*\[.*?(\d+)\s*점\]\s*([\s\S]*?)(?=\n- |\Z)")
                match = pattern.search(response_text)
                if match:
                    score = int(match.group(1))
                    reason = match.group(2).strip()
                    scores[item_name] = {"점수": score, "이유": reason, "배점": max_score}
                    total_score += score
                else:
                    scores[item_name] = {"점수": 0, "이유": "항목별 평가 추출 실패", "배점": max_score}
            parsed_data['항목별 평가'] = scores
            parsed_data['총점'] = total_score
        except Exception as e:
            return {"종합 평가": f"AI 응답 분석 실패: {e}", "항목별 평가": {}, "총점": 0}
        return parsed_data

    # --- UI 로직 ---
    st.subheader("📝 1단계: 평가 정보 입력")
    eval_name = st.text_input("평가명", placeholder="예: 2025년 1학기 중간 논술 평가")
    eval_date = st.date_input("평가일자", datetime.date.today())

    with st.expander("📊 2단계: 평가 기준 설정", expanded=True):
        if 'criteria_list' not in st.session_state:
            st.session_state.criteria_list = [{"항목": "내용의 충실성", "배점": 40, "기준": "주제에 대한 이해가 깊고, 근거가 타당하며 내용이 풍부한가?"},
                                             {"항목": "논리적 구조", "배점": 30, "기준": "서론, 본론, 결론의 구조가 명확하고, 문단 간의 연결이 자연스러운가?"},
                                             {"항목": "표현의 정확성", "배점": 30, "기준": "어휘 사용이 적절하고, 문법 및 맞춤법 오류가 없는가?"}]
        for i, criterion in enumerate(st.session_state.criteria_list):
            st.markdown("---")
            col1, col2, col3 = st.columns([3, 1, 1])
            criterion['항목'] = col1.text_input(f"항목 #{i+1}", value=criterion['항목'], key=f"item_{i}")
            criterion['배점'] = col2.number_input(f"배점 #{i+1}", min_value=0, max_value=100, value=criterion['배점'], key=f"score_{i}")
            with col3:
                st.write(""); st.write("")
                if st.button("🗑️ 삭제", key=f"delete_{i}"):
                    st.session_state.criteria_list.pop(i)
                    st.rerun()
            criterion['기준'] = st.text_area(f"세부 기준 #{i+1}", value=criterion['기준'], key=f"desc_{i}", height=100)
        if st.button("➕ 평가 항목 추가"):
            st.session_state.criteria_list.append({"항목": "", "배점": 10, "기준": ""})
            st.rerun()

    st.subheader("📄 3단계: 에세이 파일 업로드 및 평가 실행")
    uploaded_essays = st.file_uploader("평가할 학생들의 에세이 PDF 파일들을 업로드하세요.", type=['pdf'], accept_multiple_files=True)

    if st.button("🚀 모든 파일 평가 시작", key="start_eval_button"):
        if not eval_name:
            st.error("⚠️ 1단계에서 평가명을 입력해주세요!")
        elif not uploaded_essays:
            st.warning("⚠️ 3단계에서 평가할 에세이 파일을 먼저 업로드해주세요.")
        else:
            results = []
            progress_bar = st.progress(0, text="평가를 시작합니다...")
            for i, essay_file in enumerate(uploaded_essays):
                progress_text = f"평가 진행 중: {essay_file.name} ({i+1}/{len(uploaded_essays)})"
                progress_bar.progress((i + 1) / len(uploaded_essays), text=progress_text)
                try:
                    full_text = ""
                    with pdfplumber.open(essay_file) as pdf:
                        full_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
                    if not full_text.strip():
                        results.append({"파일명": essay_file.name, "평가결과_원본": "텍스트 추출 실패", "평가결과_분석": {}})
                        continue
                    criteria_str = "\n".join([f"- 항목: {c['항목']}, 배점: {c['배점']}, 기준: {c['기준']}" for c in st.session_state.criteria_list])
                    output_format_str = "\n".join([f"- {c['항목']}: [{c['배점']}점 만점에 00점]\n(점수 부여에 대한 구체적인 이유를 서술해주세요.)" for c in st.session_state.criteria_list])
                    prompt = f"""
                    당신은 전문적인 논술 평가관입니다. 주어진 학생의 에세이를 아래의 평가 기준에 따라 분석하고 채점해주세요.
                    **평가 기준:**\n{criteria_str}
                    **학생 에세이:**\n---\n{full_text}\n---
                    **출력 형식 (반드시 이 형식에 맞춰 응답해주세요. 각 항목은 줄바꿈으로 구분합니다.):**
                    [종합 평가]
                    (에세이 전체에 대한 강점과 개선점을 포함한 종합적인 피드백을 서술해주세요.)
                    [항목별 평가]
                    {output_format_str}
                    """
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    response = model.generate_content(prompt, request_options={'timeout': 300})
                    parsed_result = parse_ai_response(response.text, st.session_state.criteria_list)
                    results.append({"파일명": essay_file.name, "평가결과_원본": response.text, "평가결과_분석": parsed_result})
                    time.sleep(1)
                except Exception as e:
                    st.error(f"{essay_file.name} 평가 중 오류 발생: {e}")
                    results.append({"파일명": essay_file.name, "평가결과_원본": f"오류 발생: {e}", "평가결과_분석": {}})
            st.session_state['evaluation_results'] = results

    # --- 📈 4. 평가 결과 확인 및 다운로드 ---
    if 'evaluation_results' in st.session_state and st.session_state['evaluation_results']:
        st.subheader("📈 4단계: 평가 결과 확인 및 다운로드")
        results_data = st.session_state['evaluation_results']
        criteria_names = [c['항목'] for c in st.session_state.criteria_list]
        summary_data = []
        for result in results_data:
            row = {'파일명': result['파일명']}
            parsed_scores = result.get('평가결과_분석', {}).get('항목별 평가', {})
            for name in criteria_names:
                row[name] = parsed_scores.get(name, {}).get('점수', 'N/A')
            row['총점'] = result.get('평가결과_분석', {}).get('총점', 'N/A')
            summary_data.append(row)
        summary_df = pd.DataFrame(summary_data)
        st.markdown("### 📊 전체 점수 요약표")
        st.dataframe(summary_df)

        col1, col2 = st.columns(2)
        with col1:
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                summary_df.to_excel(writer, index=False, sheet_name='전체 점수 요약')
            st.download_button(label="📥 엑셀 요약표 다운로드", data=excel_buffer.getvalue(), file_name=f"{eval_name}_전체요약.xlsx", mime="application/vnd.ms-excel", key="download_excel")
        with col2:
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                for result in results_data:
                    report_docx_buffer = generate_report_docx(result, eval_name, eval_date)
                    zipf.writestr(f"{os.path.splitext(result['파일명'])[0]}_상세보고서.docx", report_docx_buffer.getvalue())
            st.download_button(label="🗂️ 모든 상세 보고서 (ZIP) 다운로드", data=zip_buffer.getvalue(), file_name=f"{eval_name}_상세보고서.zip", mime="application/zip", key="download_zip")

        st.markdown("### 📝 학생별 상세 평가")
        for result in results_data:
            with st.expander(f"📄 {result['파일명']} 상세 결과 보기"):
                # 화면 표시는 마크다운 함수를 사용하도록 수정
                st.markdown(generate_report_markdown(result))
                
                # 다운로드는 워드 함수를 사용
                report_docx_buffer = generate_report_docx(result, eval_name, eval_date)
                st.download_button(
                    label="📋 개별 보고서 다운로드 (.docx)",
                    data=report_docx_buffer.getvalue(),
                    file_name=f"{os.path.splitext(result['파일명'])[0]}_상세보고서.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{result['파일명']}"
                )
        
        if st.button("💾 현재 평가를 기록에 저장", key="save_history_button"):
            new_history_item = {
                "평가명": eval_name,
                "평가일자": eval_date.strftime("%Y-%m-%d"),
                "평가기준": st.session_state.criteria_list,
                "평가결과": results_data
            }
            history.append(new_history_item)
            save_history(history)
            st.success("현재 평가가 기록에 성공적으로 저장되었습니다!")
            st.rerun()
